"""Generate COS30018 HNRS Report as Word document."""
import json
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="2563EB"):
    """Add a styled table with colored header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # shade header
        tc_pr = cell._element.get_or_add_tcPr()
        shading = tc_pr.makeelement(qn('w:shd'), {
            qn('w:fill'): header_color, qn('w:val'): 'clear'
        })
        tc_pr.append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            # Alternate row shading
            if r_idx % 2 == 0:
                tc_pr = cell._element.get_or_add_tcPr()
                shading = tc_pr.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F1F5F9', qn('w:val'): 'clear'
                })
                tc_pr.append(shading)

    # Set column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return h


def para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p


# ── Main ─────────────────────────────────────────────────────────────────

def generate():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # ── Title Page ───────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("COS30018 - Intelligent Systems")
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run("Assignment Option B")
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph()

    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = main_title.add_run("Handwritten Number Recognition System\n(HNRS)")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    for _ in range(3):
        doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Swinburne University of Technology\nSemester: Summer 2025")
    run.font.size = Pt(12)

    doc.add_page_break()

    # ── Table of Contents ────────────────────────────────────────────────
    heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. System Architecture",
        "3. Task 1: Image Preprocessing",
        "4. Task 2: Image Segmentation",
        "5. Task 3: Machine Learning Models",
        "6. Task 4: Model Evaluation & Comparison",
        "7. GUI Implementation",
        "8. Extension: Arithmetic Expression Recognition",
        "9. Results Summary",
        "10. Conclusion",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)

    doc.add_page_break()

    # ── 1. Introduction ──────────────────────────────────────────────────
    heading(doc, "1. Introduction", level=1)

    para(doc,
        "This report presents the implementation of a Handwritten Number Recognition System (HNRS) "
        "for the COS30018 Intelligent Systems unit at Swinburne University of Technology. "
        "The system recognizes handwritten digits (0-9), multi-digit numbers, and full arithmetic "
        "expressions from both drawn input and uploaded images, using multiple machine learning "
        "approaches trained on the MNIST dataset."
    )

    heading(doc, "Key Features", level=2)
    bullet(doc, "Three image preprocessing techniques for noise handling")
    bullet(doc, "Three image segmentation methods for multi-digit and expression recognition")
    bullet(doc, "Three ML models: CNN (PyTorch), SVM, and KNN")
    bullet(doc, "PyQt5-based GUI with wide drawing canvas supporting single digits, multi-digit numbers, and expressions")
    bullet(doc, "Automatic segmentation and per-character classification with visual feedback")
    bullet(doc, "Extension: Arithmetic expression recognition (+, -, *, /, parentheses) with real-time evaluation")

    heading(doc, "Technology Stack", level=2)
    add_styled_table(doc,
        ["Component", "Technology", "Version"],
        [
            ["Deep Learning", "PyTorch", "2.7.1"],
            ["Traditional ML", "scikit-learn", "Latest"],
            ["Image Processing", "OpenCV", "Latest"],
            ["GUI Framework", "PyQt5", "5.x"],
            ["Visualization", "matplotlib + seaborn", "Latest"],
            ["Language", "Python", "3.13"],
        ],
        col_widths=[5, 5, 4]
    )

    doc.add_page_break()

    # ── 2. System Architecture ───────────────────────────────────────────
    heading(doc, "2. System Architecture", level=1)

    para(doc, "The system follows a modular pipeline architecture:")

    heading(doc, "Processing Pipeline", level=2)
    para(doc,
        "Input Image  -->  Preprocessing (Basic / Otsu / Adaptive)  "
        "-->  Segmentation (Contour / Connected Components / Projection)  "
        "-->  Classification (CNN / SVM / KNN)  "
        "-->  Output (Predicted digits with confidence scores)"
    )

    heading(doc, "Project Structure", level=2)
    structure_items = [
        ("main.py", "Entry point - GUI mode or --train mode"),
        ("config.py", "Global configuration & hyperparameters"),
        ("preprocessing/preprocessor.py", "3 preprocessing techniques"),
        ("segmentation/segmenter.py", "3 segmentation methods"),
        ("models/base_model.py", "Abstract base class for all models"),
        ("models/cnn_pytorch.py", "CNN implementation (PyTorch)"),
        ("models/svm_model.py", "SVM classifier (scikit-learn)"),
        ("models/knn_model.py", "KNN classifier (scikit-learn)"),
        ("models/model_manager.py", "Model lifecycle management"),
        ("evaluation/evaluator.py", "Metrics computation & visualization"),
        ("gui/", "PyQt5 GUI with 3 tabs"),
        ("extension/", "Arithmetic expression recognition"),
        ("utils/image_utils.py", "Image I/O & conversion utilities"),
    ]
    for path, desc in structure_items:
        bullet(doc, f"{path} - {desc}")

    heading(doc, "Configuration (config.py)", level=2)
    add_styled_table(doc,
        ["Parameter", "Value", "Description"],
        [
            ["IMAGE_SIZE", "28", "MNIST standard (28x28 pixels)"],
            ["BATCH_SIZE", "64", "Training batch size"],
            ["EPOCHS", "10", "Training epochs for CNN"],
            ["LEARNING_RATE", "0.001", "Adam optimizer learning rate"],
            ["VALIDATION_SPLIT", "0.1", "10% validation data"],
            ["SVM_KERNEL", "rbf", "RBF kernel for SVM"],
            ["SVM_C", "10", "SVM regularization parameter"],
            ["KNN_N_NEIGHBORS", "5", "Number of neighbors for KNN"],
        ],
        col_widths=[4.5, 3, 7]
    )

    doc.add_page_break()

    # ── 3. Task 1: Preprocessing ─────────────────────────────────────────
    heading(doc, "3. Task 1: Image Preprocessing", level=1)

    para(doc,
        "Three preprocessing techniques were implemented to handle varying input quality. "
        "Each technique converts raw input images into normalized 28x28 grayscale format "
        "suitable for the ML models."
    )

    heading(doc, "3.1 Basic Preprocessing", level=2)
    para(doc, "The simplest pipeline for clean input images:", bold=True)
    bullet(doc, "Convert to grayscale (if not already)")
    bullet(doc, "Resize to 28x28 pixels using INTER_AREA interpolation")
    bullet(doc, "Normalize pixel values to [0, 1] range")
    para(doc, "Best suited for clean, high-contrast input images.", italic=True)

    heading(doc, "3.2 Otsu Binarization", level=2)
    para(doc, "Automatic threshold selection for varying lighting conditions:", bold=True)
    bullet(doc, "Convert to grayscale")
    bullet(doc, "Apply Otsu's method (cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU) to automatically determine the optimal binarization threshold")
    bullet(doc, "Resize to 28x28 pixels")
    bullet(doc, "Normalize to [0, 1]")
    para(doc,
        "Otsu's method minimizes intra-class variance between foreground and background pixels, "
        "making it effective for images with bimodal intensity distributions.",
        italic=True
    )

    heading(doc, "3.3 Adaptive Threshold with Denoising", level=2)
    para(doc, "Most robust technique for noisy or unevenly lit images:", bold=True)
    bullet(doc, "Convert to grayscale")
    bullet(doc, "Apply Gaussian blur (5x5 kernel) for noise reduction")
    bullet(doc, "Apply adaptive threshold (Gaussian method, block size 11, constant 2) - threshold varies per-pixel based on local neighborhood")
    bullet(doc, "Apply morphological closing (3x3 rectangular kernel) to fill small gaps in digit strokes")
    bullet(doc, "Resize to 28x28 pixels")
    bullet(doc, "Normalize to [0, 1]")

    heading(doc, "3.4 Additional Utilities", level=2)
    bullet(doc, "invert_if_needed(): Automatically inverts images if mean pixel value > 127, ensuring white-on-black MNIST format")
    bullet(doc, "normalize_segmented(): Processes already-segmented digit images for model input")
    bullet(doc, "canvas_to_mnist_format(): Converts PyQt5 canvas drawings to MNIST-compatible format with bounding box extraction and 20% padding")

    doc.add_page_break()

    # ── 4. Task 2: Segmentation ──────────────────────────────────────────
    heading(doc, "4. Task 2: Image Segmentation", level=1)

    para(doc,
        "Three segmentation techniques were implemented to extract individual digits "
        "from multi-digit images. All techniques share common helper functions for "
        "binary image preparation and digit extraction with padding."
    )

    heading(doc, "4.1 Contour-Based Segmentation", level=2)
    para(doc, "Uses OpenCV's contour detection algorithm:", bold=True)
    bullet(doc, "Prepare binary image (grayscale, invert if needed, Otsu binarization)")
    bullet(doc, "Find external contours using cv2.findContours with RETR_EXTERNAL mode")
    bullet(doc, "Compute bounding boxes for each contour")
    bullet(doc, "Filter by minimum area (>1% of image area) and minimum dimensions (height >15% or width >3%)")
    bullet(doc, "Sort bounding boxes left-to-right by x-coordinate")
    bullet(doc, "Extract each digit region, pad to square, resize to 28x28")

    heading(doc, "4.2 Connected Components Segmentation", level=2)
    para(doc, "Uses pixel connectivity analysis:", bold=True)
    bullet(doc, "Prepare binary image")
    bullet(doc, "Apply cv2.connectedComponentsWithStats with 8-connectivity")
    bullet(doc, "Filter components by area (>0.5% of image) and dimension thresholds")
    bullet(doc, "Merge overlapping bounding boxes (overlap threshold 0.3) to handle digits with disconnected strokes")
    bullet(doc, "Sort left-to-right, extract and pad each digit to 28x28")

    heading(doc, "4.3 Vertical Projection Segmentation", level=2)
    para(doc, "Uses column-wise pixel density analysis:", bold=True)
    bullet(doc, "Prepare binary image")
    bullet(doc, "Compute vertical projection (sum of white pixels per column)")
    bullet(doc, "Apply threshold at max(2, 8% of maximum projection value)")
    bullet(doc, "Identify contiguous regions above threshold as digit locations")
    bullet(doc, "Detect gaps between regions as digit separators")
    bullet(doc, "Filter out segments narrower than 2 pixels or shorter than 10% of image height")

    heading(doc, "4.4 Comparison", level=2)
    add_styled_table(doc,
        ["Method", "Strengths", "Weaknesses"],
        [
            ["Contour-based", "Fast, handles clear digits well", "May miss digits with broken strokes"],
            ["Connected Components", "Handles disconnected strokes", "Slightly more complex processing"],
            ["Vertical Projection", "Good for evenly spaced digits", "Struggles with overlapping digits"],
        ],
        col_widths=[4.5, 5, 5]
    )

    doc.add_page_break()

    # ── 5. Task 3: ML Models ─────────────────────────────────────────────
    heading(doc, "5. Task 3: Machine Learning Models", level=1)

    para(doc,
        "Three machine learning models were implemented, all inheriting from an abstract "
        "BaseModel class that enforces a consistent interface: build(), train(), predict(), "
        "predict_proba(), save(), and load()."
    )

    heading(doc, "5.1 CNN (PyTorch) - Primary Model", level=2)
    para(doc, "A Convolutional Neural Network implemented in PyTorch:")

    add_styled_table(doc,
        ["Layer", "Type", "Details", "Output Shape"],
        [
            ["1", "Conv2d", "32 filters, 3x3 kernel", "(32, 26, 26)"],
            ["2", "ReLU + MaxPool2d", "2x2 pooling", "(32, 13, 13)"],
            ["3", "Conv2d", "64 filters, 3x3 kernel", "(64, 11, 11)"],
            ["4", "ReLU + MaxPool2d", "2x2 pooling", "(64, 5, 5)"],
            ["5", "Conv2d", "64 filters, 3x3 kernel", "(64, 3, 3)"],
            ["6", "ReLU", "-", "(64, 3, 3)"],
            ["7", "Flatten", "-", "(576,)"],
            ["8", "Linear + ReLU", "128 neurons", "(128,)"],
            ["9", "Dropout", "p=0.5", "(128,)"],
            ["10", "Linear (Output)", "10 classes", "(10,)"],
        ],
        col_widths=[1.5, 4, 4, 4]
    )

    doc.add_paragraph()
    para(doc, "Training Configuration:", bold=True)
    bullet(doc, "Optimizer: Adam (learning rate = 0.001)")
    bullet(doc, "Loss Function: CrossEntropyLoss")
    bullet(doc, "Batch Size: 64")
    bullet(doc, "Epochs: 10")
    bullet(doc, "Validation Split: 10%")
    bullet(doc, "Device: CPU (CUDA auto-detected if available)")
    p = para(doc, "Result: 99.28% accuracy on MNIST test set", bold=True)
    p.runs[0].font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

    heading(doc, "5.2 SVM (Support Vector Machine)", level=2)
    para(doc, "Configuration:", bold=True)
    bullet(doc, "Algorithm: sklearn.svm.SVC")
    bullet(doc, "Kernel: RBF (Radial Basis Function)")
    bullet(doc, "C = 10 (regularization parameter)")
    bullet(doc, "gamma = 'scale' (1 / (n_features * variance))")
    bullet(doc, "Probability estimation: enabled")
    bullet(doc, "Input: Flattened 28x28 = 784 features, standardized via StandardScaler")
    bullet(doc, "Training set: 20,000 samples (SVM training is O(n^2) to O(n^3))")
    p = para(doc, "Result: 95.79% accuracy on MNIST test set", bold=True)
    p.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    heading(doc, "5.3 KNN (K-Nearest Neighbors)", level=2)
    para(doc, "Configuration:", bold=True)
    bullet(doc, "Algorithm: sklearn.neighbors.KNeighborsClassifier")
    bullet(doc, "K = 5 neighbors")
    bullet(doc, "Weights: 'distance' (closer neighbors have more influence)")
    bullet(doc, "Algorithm selection: auto (ball_tree, kd_tree, or brute)")
    bullet(doc, "Parallel processing: all CPU cores (n_jobs=-1)")
    bullet(doc, "Input: Flattened 784 features, standardized via StandardScaler")
    bullet(doc, "Training set: 20,000 samples")
    p = para(doc, "Result: 93.18% accuracy on MNIST test set", bold=True)
    p.runs[0].font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

    doc.add_page_break()

    # ── 6. Task 4: Evaluation ────────────────────────────────────────────
    heading(doc, "6. Task 4: Model Evaluation & Comparison", level=1)

    # Load actual results
    results_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "data", "evaluation_results", "results.json"
    )
    with open(results_path, "r") as f:
        results = json.load(f)

    heading(doc, "6.1 Overall Accuracy Comparison", level=2)
    add_styled_table(doc,
        ["Model", "Accuracy", "Total Inference (10K)", "Avg/Sample"],
        [
            ["CNN (PyTorch)", "99.28%", "0.58s", "0.058ms"],
            ["SVM", "95.79%", "28.49s", "2.849ms"],
            ["KNN", "93.18%", "3.41s", "0.341ms"],
        ],
        col_widths=[4, 3, 4, 3]
    )

    # Per-class metrics for each model
    model_names_display = {
        "CNN (PyTorch)": "CNN (PyTorch) - Best Model",
        "SVM (scikit-learn)": "SVM",
        "KNN (scikit-learn)": "KNN"
    }

    for idx, (section_num, result) in enumerate(zip(["6.2", "6.3", "6.4"], results)):
        model_display = model_names_display.get(result["model_name"], result["model_name"])
        heading(doc, f"{section_num} Per-Class Metrics - {model_display}", level=2)

        per_class = result.get("per_class", {})
        rows = []
        for digit in range(10):
            d = per_class.get(str(digit), {})
            rows.append([
                str(digit),
                f"{d.get('precision', 0):.4f}",
                f"{d.get('recall', 0):.4f}",
                f"{d.get('f1', 0):.4f}",
                str(d.get('support', 0))
            ])

        add_styled_table(doc,
            ["Digit", "Precision", "Recall", "F1-Score", "Support"],
            rows,
            col_widths=[2, 3, 3, 3, 2.5]
        )
        doc.add_paragraph()

    heading(doc, "6.5 Analysis", level=2)

    para(doc, "CNN (PyTorch) - Best Model:", bold=True)
    bullet(doc, "Highest accuracy (99.28%) across all digits with F1 > 0.98 for every class")
    bullet(doc, "Fastest inference time (0.058ms per sample)")
    bullet(doc, "Digit '1' achieved perfect recall (1.0000)")
    bullet(doc, "Weakest digit was '9' (recall 0.9812), primarily confused with digits 4 and 7")

    para(doc, "SVM:", bold=True)
    bullet(doc, "Solid performance (95.79%) but slowest inference (2.849ms/sample) due to RBF kernel computation")
    bullet(doc, "Most confusion in digits 7, 8, and 9")
    bullet(doc, "Digit '7' had the lowest precision (0.9157)")

    para(doc, "KNN:", bold=True)
    bullet(doc, "Lowest accuracy (93.18%) but simple and fast to train")
    bullet(doc, "Most confusion in digits 2, 3, 8, and 9")
    bullet(doc, "Digit '9' had the lowest precision (0.8843)")

    para(doc,
        "Key Insight: Deep learning (CNN) significantly outperforms traditional ML methods (SVM, KNN) "
        "on image classification tasks due to its ability to learn hierarchical spatial features "
        "through convolutional layers, while traditional methods operate on flattened pixel vectors "
        "and lose spatial information.",
        italic=True
    )

    doc.add_page_break()

    # ── 7. GUI ───────────────────────────────────────────────────────────
    heading(doc, "7. GUI Implementation", level=1)

    para(doc,
        "The GUI is built with PyQt5 using a clean, light theme inspired by modern web design. "
        "It features a three-tab interface for different functionalities."
    )

    heading(doc, "7.1 Tab 1: Recognition", level=2)
    para(doc,
        "The Recognition tab is the primary interface for handwritten input recognition. "
        "It supports single digits, multi-digit numbers, and full arithmetic expressions."
    )
    bullet(doc, "Wide drawing canvas (560x200 pixels) for writing expressions like '23+45' or '9*8-1'")
    bullet(doc, "Automatic segmentation: the system segments drawn input into individual characters, then classifies each one")
    bullet(doc, "Expression Mode (enabled by default): uses the 16-class ExpressionCNN to recognize digits AND operators (+, -, *, /), then evaluates the expression and shows the computed result")
    bullet(doc, "Digit-only Mode: when Expression Mode is off, uses the selected digit model (CNN/SVM/KNN) for digit-only recognition of multi-digit numbers")
    bullet(doc, "Model selector: CNN (PyTorch), SVM, KNN - switchable at runtime")
    bullet(doc, "Preprocessing method selector: Basic, Otsu Binarization, Adaptive Threshold")
    bullet(doc, "Segmentation method selector: Contour-based, Connected Components, Vertical Projection")
    bullet(doc, "Image upload support (single file or entire folder of digit images)")
    bullet(doc, "Segmented Symbols panel: displays each segmented character image alongside its predicted label, allowing the user to verify the segmentation and classification quality")
    bullet(doc, "Result display: shows the recognized expression/number in large text, with computed result (= answer) shown in a green highlight box for expressions")
    bullet(doc, "Confidence distribution bar chart showing probabilities for all 10 digit classes")

    heading(doc, "7.2 Tab 2: Evaluation", level=2)
    bullet(doc, "Summary cards showing accuracy and inference time for each model")
    bullet(doc, "Accuracy comparison bar chart (matplotlib)")
    bullet(doc, "Confusion matrix heatmap (seaborn) with navigation between models")
    bullet(doc, "Per-class metrics table (Precision, Recall, F1-Score)")
    bullet(doc, "All results loaded from pre-computed JSON file for instant display")

    heading(doc, "7.3 Tab 3: Models", level=2)
    bullet(doc, "Model information cards with architecture details, accuracy, and training time")
    bullet(doc, "Status indicators showing whether each model is ready or needs training")
    bullet(doc, "Optional re-training section with configurable hyperparameters (epochs, batch size)")
    bullet(doc, "Training progress bar and real-time log output")

    heading(doc, "7.4 Design", level=2)
    bullet(doc, "Clean light theme: white card backgrounds (#ffffff) with subtle borders (#e2e8f0)")
    bullet(doc, "Blue primary color (#2563eb) for highlights and interactive elements")
    bullet(doc, "Green success indicators (#16a34a) for positive status")
    bullet(doc, "Segoe UI font family throughout for consistent readability")
    bullet(doc, "Rounded corners and consistent spacing for modern appearance")

    doc.add_page_break()

    # ── 8. Extension ─────────────────────────────────────────────────────
    heading(doc, "8. Extension: Arithmetic Expression Recognition", level=1)

    para(doc,
        "The extension adds the ability to recognize and evaluate handwritten arithmetic "
        "expressions containing digits (0-9) and operators (+, -, *, /, parentheses)."
    )

    heading(doc, "8.1 16-Class Expression CNN", level=2)
    para(doc, "The digit CNN architecture was extended to classify 16 symbols:")

    add_styled_table(doc,
        ["Class", "Symbol", "Class", "Symbol"],
        [
            ["0-9", "Digits 0 through 9", "10", "+ (addition)"],
            ["11", "- (subtraction)", "12", "* (multiplication)"],
            ["13", "/ (division)", "14", "( (left parenthesis)"],
            ["15", ") (right parenthesis)", "", ""],
        ],
        col_widths=[2.5, 5, 2.5, 5]
    )

    para(doc,
        "Architecture: Same CNN structure as the digit model (3 Conv layers + 2 Dense layers) "
        "but with 16 output neurons instead of 10."
    )

    heading(doc, "8.2 Synthetic Training Data", level=2)
    para(doc, "Since no standard dataset exists for handwritten math operators:")
    bullet(doc, "3,000 synthetic images generated per operator class using OpenCV drawing functions")
    bullet(doc, "Random variations in position, size, and stroke thickness for diversity")
    bullet(doc, "Data augmentation: Gaussian noise (0.05 level) + random rotation (-10 to +10 degrees)")
    bullet(doc, "Combined with 3,000 MNIST samples per digit class (balanced)")
    bullet(doc, "Total training set: ~48,000 samples (30,000 digit + 18,000 operator)")

    heading(doc, "8.3 Expression Recognition Pipeline", level=2)
    bullet(doc, "Step 1: Segment the expression image into individual symbols using selected segmentation method")
    bullet(doc, "Step 2: Classify each symbol using the 16-class ExpressionCNN")
    bullet(doc, "Step 3: Build expression string - concatenate consecutive digits into multi-digit numbers, insert operators")
    bullet(doc, "Step 4: Safe evaluation using Python eval() with restricted builtins (no access to built-in functions)")
    bullet(doc, "Step 5: Error handling for ZeroDivisionError, SyntaxError, and invalid expressions")

    heading(doc, "8.4 Safety Measures", level=2)
    bullet(doc, "Whitelist of allowed characters: 0-9, +, -, *, /, (, ), decimal point, space")
    bullet(doc, "Restricted builtins dictionary: {'__builtins__': {}} prevents code injection")
    bullet(doc, "Float results rounded to 6 decimal places")
    bullet(doc, "Graceful error messages for invalid expressions")

    doc.add_page_break()

    # ── 9. Results Summary ───────────────────────────────────────────────
    heading(doc, "9. Results Summary", level=1)

    heading(doc, "9.1 Assignment Requirements Checklist", level=2)
    add_styled_table(doc,
        ["Requirement", "Marks", "Status", "Details"],
        [
            ["Task 1: Preprocessing", "8", "Complete", "3 techniques: Basic, Otsu, Adaptive"],
            ["Task 2: Segmentation", "8", "Complete", "3 techniques: Contour, Connected Components, Projection"],
            ["Task 3: ML Models", "24", "Complete", "CNN (99.28%), SVM (95.79%), KNN (93.18%)"],
            ["Task 4: Evaluation", "10", "Complete", "Accuracy, Precision, Recall, F1, Confusion Matrix"],
            ["GUI", "10", "Complete", "PyQt5 with 3 tabs, canvas, image upload"],
            ["Extension: Expression", "20", "Complete", "16-class CNN, operator recognition, evaluation"],
            ["Good Programming", "-", "Complete", "Modular design, abstract base, documented code"],
        ],
        col_widths=[4, 1.5, 2, 7]
    )

    heading(doc, "9.2 Model Performance Summary", level=2)
    add_styled_table(doc,
        ["Metric", "CNN (PyTorch)", "SVM", "KNN"],
        [
            ["Accuracy", "99.28%", "95.79%", "93.18%"],
            ["Inference/sample", "0.058ms", "2.849ms", "0.341ms"],
            ["Training time", "~158.8s", "~211.4s", "~2.5s"],
            ["Architecture", "3 Conv + 2 Dense", "RBF, C=10", "k=5, distance"],
            ["Input format", "28x28 image", "784 vector", "784 vector"],
            ["Save format", ".pth", ".pkl", ".pkl"],
        ],
        col_widths=[3.5, 3.5, 3.5, 3.5]
    )

    doc.add_page_break()

    # ── 10. Conclusion ───────────────────────────────────────────────────
    heading(doc, "10. Conclusion", level=1)

    para(doc,
        "This project successfully implements a complete Handwritten Number Recognition System "
        "with multiple machine learning approaches. The system demonstrates the full pipeline "
        "from raw image input to predicted output, with comprehensive evaluation and a "
        "user-friendly graphical interface."
    )

    heading(doc, "Key Findings", level=2)

    bullet(doc,
        "CNN (PyTorch) achieves the highest accuracy (99.28%) and fastest inference time, "
        "demonstrating the superiority of deep learning for image classification tasks."
    )
    bullet(doc,
        "SVM provides solid performance (95.79%) with RBF kernel but has significantly "
        "slower inference due to kernel computation on high-dimensional data."
    )
    bullet(doc,
        "KNN offers the simplest approach with acceptable accuracy (93.18%) and extremely "
        "fast training time (~2.5s), making it suitable for rapid prototyping."
    )
    bullet(doc,
        "The modular architecture with abstract base class enables easy addition of new models "
        "and ensures consistent evaluation across all approaches."
    )
    bullet(doc,
        "The arithmetic expression recognition extension demonstrates practical application "
        "beyond single-digit classification, combining segmentation, multi-class classification, "
        "and expression parsing into a complete pipeline."
    )

    para(doc,
        "The system provides a comprehensive solution for handwritten digit and expression "
        "recognition, with clear comparisons between different ML approaches and practical "
        "tools for interactive use through the PyQt5 GUI.",
        italic=True
    )

    # ── Footer ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("COS30018 - Intelligent Systems | Swinburne University of Technology")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run.italic = True

    # ── Save ─────────────────────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "COS30018_HNRS_Report.docx")
    doc.save(output_path)
    print(f"Report saved to: {output_path}")


if __name__ == "__main__":
    generate()
